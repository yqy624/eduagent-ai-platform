"""文档加载器"""
import os
from typing import List, Optional
from langchain_core.documents import Document


class DocumentLoader:
    """从不同格式的文件中提取文本内容"""

    @staticmethod
    def load_pdf(file_path: str) -> str:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or "")
        return "\n".join(text)

    @staticmethod
    def load_docx(file_path: str) -> str:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        return "\n".join([p.text for p in doc.paragraphs])

    @staticmethod
    def load_txt(file_path: str, encoding: str = "utf-8") -> str:
        with open(file_path, "r", encoding=encoding) as f:
            return f.read()

    @staticmethod
    def load_markdown(file_path: str) -> str:
        return DocumentLoader.load_txt(file_path)

    @staticmethod
    def load(file_path: str) -> Optional[str]:
        ext = os.path.splitext(file_path)[1].lower()
        loaders = {
            ".pdf": DocumentLoader.load_pdf,
            ".docx": DocumentLoader.load_docx,
            ".doc": lambda p: DocumentLoader.load_docx(p),
            ".txt": DocumentLoader.load_txt,
            ".md": DocumentLoader.load_markdown,
        }
        loader = loaders.get(ext)
        if loader is None:
            return None
        try:
            return loader(file_path)
        except Exception as e:
            print(f"加载文件失败 {file_path}: {e}")
            return None


class TextSplitter:
    """文本切片"""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_text(self, text: str) -> List[str]:
        """按字符长度切片，保留段落边界"""
        if not text:
            return []
        paragraphs = text.split("\n")
        chunks = []
        current = []
        current_len = 0
        for para in paragraphs:
            para_len = len(para)
            if current_len + para_len > self.chunk_size and current:
                chunks.append("\n".join(current))
                # overlap: 保留最后一小段
                overlap_text = "\n".join(current[-2:]) if len(current) >= 2 else "\n".join(current)
                current = [overlap_text] if len(overlap_text) < self.chunk_overlap * 2 else []
                current_len = sum(len(p) for p in current)
            current.append(para)
            current_len += para_len
        if current:
            chunks.append("\n".join(current))
        return chunks

    def split_documents(self, documents: List[Document]) -> List[Document]:
        """切片文档对象"""
        result = []
        for doc in documents:
            chunks = self.split_text(doc.page_content)
            for i, chunk in enumerate(chunks):
                metadata = dict(doc.metadata)
                metadata["chunk_index"] = i
                result.append(Document(page_content=chunk, metadata=metadata))
        return result
